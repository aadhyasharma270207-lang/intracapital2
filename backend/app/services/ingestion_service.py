import os
import io
import json
import logging
import re
from typing import List, Dict, Any, Tuple
import pandas as pd
from docx import Document as DocxDocument

# Optional PyMuPDF (fitz) import
HAS_FITZ = False
try:
    import fitz # PyMuPDF
    HAS_FITZ = True
except ImportError:
    pass

# Optional Docling import
HAS_DOCLING = False
try:
    from docling.document_converter import DocumentConverter
    HAS_DOCLING = True
except ImportError:
    pass

logger = logging.getLogger(__name__)

class IngestionService:
    @classmethod
    def parse_file(cls, file_path: str, file_name: str) -> Dict[str, Any]:
        """
        Parses a file path and extracts its textual content, tabular structures, and metadata.
        Supports PDF, DOCX, CSV, XLSX, JSON, TXT.
        """
        ext = os.path.splitext(file_name)[1].lower()
        content = ""
        tables_summary = []
        metadata = {
            "file_name": file_name,
            "file_size_bytes": os.path.getsize(file_path),
            "extension": ext
        }

        try:
            if ext == ".txt":
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()
                    
            elif ext == ".json":
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    data = json.load(f)
                    content = json.dumps(data, indent=2)
                    metadata["json_keys"] = list(data.keys()) if isinstance(data, dict) else []

            elif ext in [".csv", ".xlsx"]:
                content, df_meta = cls._parse_table_file(file_path, ext)
                metadata.update(df_meta)

            elif ext == ".docx":
                content = cls._parse_docx(file_path)

            elif ext == ".pdf":
                content = cls._parse_pdf(file_path)
                
            else:
                # Basic binary/text fallback
                try:
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        content = f.read()
                except Exception:
                    content = f"[Binary File: {file_name}]"
                    
        except Exception as e:
            logger.error(f"Error parsing file {file_name}: {str(e)}")
            content = f"Parsing Error: {str(e)}"
            metadata["parsing_error"] = str(e)

        # Extract entities / metadata signals
        signals = cls._extract_metadata_signals(content)
        metadata.update(signals)

        return {
            "content": content,
            "metadata": metadata
        }

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """
        Extract text from Microsoft Word documents.
        """
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Pull table contents if any
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    tables_text.append(row_text)
                    
        return "\n".join(paragraphs + tables_text)

    @staticmethod
    def _parse_table_file(file_path: str, ext: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse CSV or Excel using pandas and generate a textual summary of columns and rows.
        """
        if ext == ".csv":
            df = pd.read_csv(file_path, encoding="utf-8", errors="ignore")
        else:
            df = pd.read_excel(file_path)
            
        columns = list(df.columns)
        num_rows = len(df)
        
        # Select first 10 rows to show as sample data in content representation
        sample_rows = df.head(15).to_string()
        
        text_content = (
            f"Table Data Summary for {os.path.basename(file_path)}:\n"
            f"Columns: {', '.join(columns)}\n"
            f"Total Row Count: {num_rows}\n"
            f"Sample Data:\n{sample_rows}"
        )
        
        metadata = {
            "columns": columns,
            "row_count": num_rows,
            "has_numeric_columns": any(pd.api.types.is_numeric_dtype(df[col]) for col in df.columns)
        }
        
        return text_content, metadata

    @classmethod
    def _parse_pdf(cls, file_path: str) -> str:
        """
        Extract text from PDF. Tries Docling first, then PyMuPDF (fitz), then a basic parser.
        """
        if HAS_DOCLING:
            try:
                logger.info("Using Docling converter for PDF parsing.")
                converter = DocumentConverter()
                result = converter.convert(file_path)
                return result.document.export_to_markdown()
            except Exception as e:
                logger.warning(f"Docling conversion failed: {str(e)}. Falling back to PyMuPDF.")

        if HAS_FITZ:
            try:
                logger.info("Using PyMuPDF (fitz) for PDF text extraction.")
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text() + "\n"
                return text
            except Exception as e:
                logger.error(f"PyMuPDF parsing failed: {str(e)}")
                
        # Basic text fallback if other tools are missing
        return f"[PDF parsing could not be completed. Install pymupdf or docling. Target: {os.path.basename(file_path)}]"

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Splits text into chunks of roughly `chunk_size` characters, with `overlap` character buffer.
        """
        if not text:
            return []
            
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start += chunk_size - overlap
            
        return chunks

    @staticmethod
    def _extract_metadata_signals(content: str) -> Dict[str, Any]:
        """
        Performs regular expression parsing on text content to extract keywords, dates,
        and potential departments/technologies as system signals.
        """
        signals = {
            "departments": [],
            "technologies": [],
            "dates": [],
            "operational_signals": []
        }
        
        # Simple entity definitions
        dept_keywords = ["Logistics", "Operations", "Sales", "Customer Service", "IT", "Engineering", "QA", "R&D", "Legal", "Finance"]
        tech_keywords = ["IoT", "Sensor", "API", "Database", "SaaS", "Telemetry", "HVAC", "Refrigeration", "AI", "Machine Learning", "Cloud", "REST"]
        signal_keywords = ["deviation", "failure", "spoilage", "anomaly", "claim", "complaint", "repair", "breakdown", "alert", "error"]

        content_lower = content.lower()
        
        for dept in dept_keywords:
            if dept.lower() in content_lower:
                signals["departments"].append(dept)
                
        for tech in tech_keywords:
            if tech.lower() in content_lower:
                signals["technologies"].append(tech)
                
        for sig in signal_keywords:
            if sig.lower() in content_lower:
                signals["operational_signals"].append(sig)

        # Basic date extraction regex
        date_pattern = r'\b(19|20)\d{2}[-/](0[1-9]|1[0-2])[-/](0[1-9]|[12]\d|3[01])\b'
        dates_found = re.findall(date_pattern, content)
        if dates_found:
            signals["dates"] = ["-".join(d) for d in dates_found[:5]]

        return signals
